from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn  # type: ignore
from docx.oxml import OxmlElement

#               Docx Generator
def _heading(doc, text, color="1A56DB"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(14)
    hr,hg,hb = int(color[:2],16),int(color[2:4],16),int(color[4:],16)
    r.font.color.rgb = RGBColor(hr,hg,hb)
    pPr = p._p.get_or_add_pPr(); pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom"); bot.set(qn("w:val"),"single")
    bot.set(qn("w:sz"),"6"); bot.set(qn("w:space"),"4"); bot.set(qn("w:color"),color)
    pBdr.append(bot); pPr.append(pBdr)

def _tbl(doc, rows):
    t = doc.add_table(rows=0, cols=2); t.style = "Table Grid"
    t.columns[0].width = Inches(1.5); t.columns[1].width = Inches(5.5)
    for label, value, *rest in rows:
        if not value or str(value).strip() in ("","-","N/A"): continue
        row = t.add_row()
        lp = row.cells[0].paragraphs[0]; lr = lp.add_run(label)
        lr.bold=True; lr.font.size=Pt(9.5); lr.font.color.rgb=RGBColor(0x37,0x41,0x51)
        vp = row.cells[1].paragraphs[0]; vr = vp.add_run(str(value)); vr.font.size=Pt(9.5)
        if rest and rest[0]: vr.font.color.rgb = rest[0]

def save_docx(internships, filters, mode, platforms_used, output_path):
    doc = Document()
    for s in doc.sections:
        s.top_margin=Inches(0.9); s.bottom_margin=Inches(0.9)
        s.left_margin=Inches(1.0); s.right_margin=Inches(1.0)
    # Title
    tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = tp.add_run("Internship Search Results"); tr.bold=True; tr.font.size=Pt(20)
    tr.font.color.rgb = RGBColor(0x11,0x18,0x27)
    sp = doc.add_paragraph(); sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp.add_run(
        f"Generated: {datetime.now().strftime('%d %B %Y')}  |  "
        f"Mode: {'0 — Auto' if mode==0 else '1 — Custom'}  |  "
        f"Platforms: {', '.join(platforms_used)}  |  {len(internships)} results"
    ).font.size = Pt(9.5)
    doc.add_paragraph()
    # Filters used
    _heading(doc, "Active Filters")
    _tbl(doc, [
        ("Mode",       f"{'0 — All default platforms' if mode==0 else '1 — Custom selection'}"),
        ("Platforms",  "  ·  ".join(platforms_used)),
        ("Keywords",   "  ·  ".join(filters.keywords) if filters.keywords else "Any"),
        ("Domain",     filters.domain or "Any"),
        ("Work Mode",  filters.mode or "Any"),
        ("Min Stipend",f"INR {filters.min_stipend:,}/month" if filters.min_stipend else "Any"),
        ("Location",   filters.location or "Any"),
    ])
    doc.add_paragraph()
    # Results by platform
    grouped = {}
    for i in internships: grouped.setdefault(i.platform, []).append(i)
    for plat, items in grouped.items():
        _heading(doc, f"{plat}  ({len(items)} results)")
        for idx, i in enumerate(items, 1):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            r = p.add_run(f"{idx}. {i.title}"); r.bold=True; r.font.size=Pt(12)
            r.font.color.rgb = RGBColor(0x11,0x18,0x27)
            _tbl(doc, [
                ("Company",  i.company),
                ("Apply",    i.url,      RGBColor(0x1A,0x56,0xDB)),
                ("Location", i.location),
                ("Mode",     i.mode),
                ("Stipend",  i.stipend,  RGBColor(0x05,0x96,0x69)),
                ("Duration", i.duration),
                ("Domain",   i.domain),
                ("Skills",   "  ·  ".join(i.skills) if i.skills else ""),
                ("Deadline", i.deadline),
                ("Posted",   i.posted),
                ("Note",     i.description[:200] if i.description else ""),
            ])
            doc.add_paragraph()
    doc.save(output_path)
    print(f"  Saved -> {output_path}")
